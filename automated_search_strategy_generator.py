"""
Automated Search Strategy Generator & Resume Analyzer
--------------------------------------------------------
This is an application that uses a T5 model to generate Boolean queries from
job descriptions, perform hybrid searches, and provide a PyQt5 GUI for resume review and export.
I wrote this during a long coding session—it's messy and could use some refactoring later.
"""

# Standard libraries 
import sys, os, csv, json, math, logging, webbrowser, re
from datetime import datetime
from urllib.parse import quote_plus

# Third-party packages 
import pandas as pd
from docx import Document
from docx.shared import Pt

from PyQt5 import QtCore, QtWidgets, QtGui

# Hugging Face stuff for T5 (I grabbed these examples from online resources, i,e huggingface)
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline, Seq2SeqTrainingArguments, Seq2SeqTrainer
import torch
from torch.utils.data import Dataset
import torch.nn.functional as F
import difflib

# Set logging level (I usually use INFO, but sometimes I want more details)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# ---------------------
# UTILITY FUNCTIONS (Messy but functional)
# ---------------------

def normalize_text(txt):
    # crude normalization: remove stuff inside parentheses, make lowercase, remove punctuation
    txt = re.sub(r'\([^)]*\)', '', txt)
    txt = txt.lower()
    return re.sub(r'[\W_]+', ' ', txt)

def estimate_years_of_experience(resumeTxt):
    # crude year range check: expecting "YYYY - YYYY" or "YYYY - present"
    patt = r'(\d{4})\s*-\s*(\d{4}|present)'
    found = re.findall(patt, resumeTxt.lower())
    if not found:
        return 0
    earliest = datetime.now().year
    latest = 0
    for start, end in found:
        try:
            s_year = int(start)
        except:
            continue  # skip if conversion fails
        if s_year < earliest:
            earliest = s_year
        e_year = int(end) if end.isdigit() else datetime.now().year
        if e_year > latest:
            latest = e_year
    return max(0, latest - earliest)

# A quick synonyms dictionary I tossed together—could be improved!
SYNONYMS = {
    "etl processes": ["etl processes", "extraction transform load", "cloud-based data analytics platforms", "data pipelines"],
    "data pipeline architectures": ["data pipeline architectures", "data pipelines", "cloud-based data analytics platforms"],
    "aws": ["aws", "amazon web services", "amazon redshift", "cloud-based data analytics platforms"],
    "gcp": ["gcp", "google cloud platform", "google cloud"],
    "azure": ["azure", "microsoft azure"]
}

def fuzzy_phrase_match(candText, phrase, word_thresh=0.5, word_sim=0.7):
    # Try exact matching first, then fallback to fuzzy matching with difflib
    cand = normalize_text(candText)
    phr = normalize_text(phrase)
    cand_words = cand.split()
    phr_words = phr.split()
    if not phr_words:
        print("Empty phrase in fuzzy_phrase_match")  # debug print
        return False
    match_count = 0
    for w in phr_words:
        if re.search(r'\b' + re.escape(w) + r'\b', cand):
            match_count += 1
        else:
            for cw in cand_words:
                if difflib.SequenceMatcher(None, w, cw).ratio() >= word_sim:
                    match_count += 1
                    break
    return (match_count / len(phr_words)) >= word_thresh

def extract_value(token):
    # Quick extraction method
    token = token.strip()
    if ":" in token:
        token = token.split(":", 1)[1].strip()
    if token.startswith('"') and token.endswith('"'):
        token = token[1:-1]
    return token

def lenient_fuzzy_ratio(a, b):
    return difflib.SequenceMatcher(None, normalize_text(a), normalize_text(b)).ratio()

def match_job_title(cand, titles):
    best = max(lenient_fuzzy_ratio(cand, t) for t in titles)
    if best >= 0.8:
        return 1.0
    elif best >= 0.6:
        return 0.9
    return 0.0

def match_location(cand, loc):
    return 1.0 if normalize_text(loc) in normalize_text(cand) else 0.9

def match_experience(cand, req):
    norm = normalize_text(cand)
    if "5+ years" in req.lower():
        yrs = estimate_years_of_experience(cand)
        if yrs >= 5:
            return 1.0
        elif yrs >= 3:
            return 0.9
        else:
            return 0.0
    if "bachelor" in req.lower():
        if "bachelor" in norm:
            return 1.0
        if "master" in norm or "mba" in norm:
            return 0.9
    if any(x in norm for x in ["computer science", "information systems", "engineering"]):
        return 1.0
    return 0.0

def match_certification(cand, certs):
    norm = normalize_text(cand)
    best = max(difflib.SequenceMatcher(None, norm, normalize_text(cert)).ratio() for cert in certs)
    return 1.0 if best >= 0.6 else 0.0

def match_cloud_platform(cand, plat_str):
    norm = normalize_text(cand)
    possibles = re.split(r',| or ', plat_str, flags=re.IGNORECASE)
    for p in possibles:
        p_norm = normalize_text(p)
        for key, syns in SYNONYMS.items():
            if p_norm in key or key in p_norm:
                if any(s in norm for s in syns):
                    return 1.0
    return 0.0

def match_skills(cand, skills):
    norm = normalize_text(cand)
    if not skills:
        return 0.0
    matched = 0
    for skill in skills:
        s_norm = normalize_text(skill)
        if any(s_norm in key or key in s_norm for key in SYNONYMS) or (s_norm in norm):
            matched += 1
    frac = matched / len(skills)
    return 1.0 if frac >= 0.5 else frac

# ---------------------
# Weighted Boolean Search 
# ---------------------
def eval_candidate_match(cand_text, query_str, job_txt=None, min_score=0.5):
    """
    Yeah, this is our main candidate scoring logic.
    TODO: Clean this mess up later.
    """
    # Strip off any trailing "AND" 
    q = query_str.strip()
    if q.endswith("AND"):
        q = q[:-3]

    # Split the query into clauses assuming "AND" is the separator
    try:
        clauses = re.split(r'\s+AND\s+', q, flags=re.IGNORECASE)
    except Exception as e:
        print("Regex splitting failed:", e)  # Ugh, debug print left in
        return False, 0.0

    n = len(clauses)
    # If exactly 6 clauses, use some custom weights 
    if n == 6:
        weights = [0.2, 0.1, 0.2, 0.2, 0.2, 0.1]
    else:
        weights = [1.0 / n for _ in range(n)]

    total_score = 0.0
    for idx, raw_clause in enumerate(clauses):
        # Lowercase and strip unnecessary parentheses (yeah, it's a bit messy)
        cl = raw_clause.strip("() ").lower()
        try:
            # Sometimes there are multiple conditions separated by OR
            options = re.split(r'\s+OR\s+', cl)
        except:
            options = [cl]
        score_list = []
        for opt in options:
            opt = opt.strip()
            if ':' in opt:
                key, val = opt.split(":", 1)
                key = key.strip().lower()
                val = val.strip()
            else:
                key, val = "misc", opt  # No key? Just treat it as misc

            # Kinda hacky matching logic here – not super scientific
            if key == "job title":
                s = match_job_title(cand_text, [val])
            elif key == "location":
                s = match_location(cand_text, val)
            elif key in ("experience", "degree", "education"):
                s = match_experience(cand_text, val)
            elif key == "certification":
                s = match_certification(cand_text, val.split(","))
            elif key == "cloud platform":
                s = match_cloud_platform(cand_text, val)
            elif key == "skills":
                s = match_skills(cand_text, [v.strip() for v in val.split(",") if v.strip()])
            else:
                # fallback fuzzy matching – works well enough
                s = 1.0 if fuzzy_phrase_match(cand_text, val, word_thresh=0.5, word_sim=0.7) else 0.0
            score_list.append(s)
        max_s = max(score_list) if score_list else 0.0
        total_score += max_s * weights[idx]

    passed = total_score >= min_score
    return passed, total_score

# ---------------------
# T5 Model Integration Functions
# ---------------------
def compute_t5_embedding(txt, model, tokenizer):
    # Just use the model's encoder to get a mean-pooled embedding
    inputs = tokenizer(txt, return_tensors="pt", truncation=True, padding=True)
    with torch.no_grad():
        output = model.encoder(**inputs)
    return output.last_hidden_state.mean(dim=1).squeeze()

def cosine_similarity_t5(emb1, emb2):
    return F.cosine_similarity(emb1, emb2, dim=0).item()

def hybrid_search_t5(cand, bool_query, query_txt, model, tokenizer, alpha=0.9, beta=0.1, thresh=0):
    # Combine Boolean score with T5 semantic similarity
    _, bool_score = eval_candidate_match(cand, bool_query, query_txt, min_score=0)
    emb_cand = compute_t5_embedding(cand, model, tokenizer)
    emb_query = compute_t5_embedding(query_txt, model, tokenizer)
    sem_score = cosine_similarity_t5(emb_cand, emb_query)
    combined = alpha * bool_score + beta * sem_score
    return combined >= thresh, combined

# ---------------------
# Dataset & Training Setup
# ---------------------
class BooleanQueryDataset(Dataset):
    def __init__(self, csv_file, tokenizer, max_inp=512, max_targ=128):
        # WARNING: Loads entire CSV into memory—might be problematic for huge files.
        self.data = pd.read_csv(csv_file)
        self.tokenizer = tokenizer
        self.max_inp = max_inp
        self.max_targ = max_targ

    def __len__(self):
        return len(self.data)

    def __getitem__(self, idx):
        job_desc = str(self.data.loc[idx, "job_description"])
        targ = str(self.data.loc[idx, "query"])
        inp_text = f"Job Description: {job_desc}\nGenerate Boolean Query:"
        inputs = self.tokenizer(inp_text, truncation=True, max_length=self.max_inp, padding="max_length", return_tensors="pt")
        targets = self.tokenizer(targ, truncation=True, max_length=self.max_targ, padding="max_length", return_tensors="pt")
        inp_ids = inputs.input_ids.squeeze()
        attn_mask = inputs.attention_mask.squeeze()
        labels = targets.input_ids.squeeze()
        labels[labels == self.tokenizer.pad_token_id] = -100
        return {"input_ids": inp_ids, "attention_mask": attn_mask, "labels": labels}

def train_model():
    # TODO: Consider adding a validation split in the future
    csv_file = "Cleaned_Dataset.csv"
    model_ckpt = "./local_t5_base"
    out_dir = "fine-tuned-t5"
    tokenizer = AutoTokenizer.from_pretrained(model_ckpt)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_ckpt)
    dataset = BooleanQueryDataset(csv_file, tokenizer)
    training_args = Seq2SeqTrainingArguments(
        output_dir=out_dir,
        evaluation_strategy="no",
        per_device_train_batch_size=4,
        num_train_epochs=3,
        save_total_limit=2,
        logging_steps=100,
        save_steps=500,
        fp16=False,
        overwrite_output_dir=True,
    )
    trainer = Seq2SeqTrainer(
        model=model,
        args=training_args,
        train_dataset=dataset,
        tokenizer=tokenizer,
    )
    trainer.train()
    trainer.save_model(out_dir)
    tokenizer.save_pretrained(out_dir)

# ---------------------
# Query Simplification Functions
# ---------------------
def simplify_query(query):
    # Balance parentheses crudely (not perfect, but it'll do for now)
    if query.count("(") > query.count(")"):
        query += ")" * (query.count("(") - query.count(")"))
    parts = []
    curr = ""
    depth = 0
    i = 0
    while i < len(query):
        ch = query[i]
        if ch == "(":
            depth += 1
            curr += ch
        elif ch == ")":
            depth -= 1
            curr += ch
        elif depth == 0 and query[i:i+5] == " AND ":
            parts.append(curr.strip())
            curr = ""
            i += 4  # Skip " AND "
        else:
            curr += ch
        i += 1
    if curr.strip():
        parts.append(curr.strip())

    def simplify_clause(clause):
        clause = clause.strip()
        inner = clause[1:-1].strip() if clause.startswith("(") and clause.endswith(")") else clause
        if "OR" in inner:
            tokens = re.split(r'\s+OR\s+', inner)
            uniq = []
            seen = set()
            for token in tokens:
                norm = re.sub(r"\s+", " ", token.strip("() ").lower())
                if norm not in seen:
                    seen.add(norm)
                    uniq.append(token.strip("() ").strip())
            return "(" + " OR ".join(uniq) + ")"
        return clause

    simp_parts = [simplify_clause(p) for p in parts]
    simp_query = " AND ".join(simp_parts)
    return re.sub(r'\s+', ' ', simp_query).strip()

def flatten_for_platform(query):
    # Remove logical operators and extra characters for external search engines
    q = query.replace(" AND ", " ")
    q = re.sub(r'[()]', '', q)
    return re.sub(r'\s+', ' ', q).strip()

# ---------------------
# Fine-Tuned Search Strategy Generator
# ---------------------
class FineTunedSearchStrategyGenerator:
    def __init__(self, model_ckpt="fine-tuned-t5"):
        # I spent way too long fine-tuning this model...
        self.tokenizer = AutoTokenizer.from_pretrained(model_ckpt)
        self.model = AutoModelForSeq2SeqLM.from_pretrained(model_ckpt)
        self.gen_pipeline = pipeline("text2text-generation",
                                     model=self.model,
                                     tokenizer=self.tokenizer,
                                     max_length=128,
                                     do_sample=False)

    def generate_boolean_query(self, job_desc):
        prompt = f"Job Description: {job_desc}\nGenerate Boolean Query:"
        print("Generating query for job description (first 50 chars):", job_desc[:50])
        generated = self.gen_pipeline(prompt, max_length=128)
        raw_query = generated[0]['generated_text']
        final_query = simplify_query(raw_query)
        return final_query, [], []

def naive_boolean_search(cand, bool_query):
    # Super simple Boolean search for fallback
    cand_lower = cand.lower()
    for part in re.split(r'\s+AND\s+', bool_query, flags=re.IGNORECASE):
        or_parts = re.split(r'\s+OR\s+', part.strip("() "), flags=re.IGNORECASE)
        if not any(re.search(r'\b' + re.escape(token.strip("() ").lower()) + r'\b', cand_lower) for token in or_parts):
            return False
    return True

# ---------------------
# GUI COMPONENTS (Messy)
# ---------------------
class FuturisticWidget(QtWidgets.QWidget):
    """A widget with an animated gradient background (just for looks, honestly)."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.phase = 0.0
        self.timer = QtCore.QTimer(self)
        self.timer.timeout.connect(self.update_phase)
        self.timer.start(50)  # Might be a bit fast—adjust as needed

    def update_phase(self):
        self.phase += 0.02
        self.update()

    def paintEvent(self, event):
        painter = QtGui.QPainter(self)
        rect = self.rect()
        gradient = QtGui.QLinearGradient(rect.topLeft(), rect.bottomRight())
        gradient.setColorAt(0, QtGui.QColor(10, 10, 30))
        mid_val = int(127 + 127 * math.sin(self.phase))
        gradient.setColorAt(0.5, QtGui.QColor(20, 20, mid_val))
        gradient.setColorAt(1, QtGui.QColor(10, 10, 30))
        painter.fillRect(rect, gradient)
        super().paintEvent(event)

class ResumeDialog(QtWidgets.QDialog):
    """A simple dialog to display a candidate's resume in HTML format."""
    def __init__(self, candName, resumeText, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Resume - {candName}")
        self.resize(600, 800)
        lay = QtWidgets.QVBoxLayout(self)
        self.text_edit = QtWidgets.QTextEdit(self)
        self.text_edit.setReadOnly(True)
        lines = resumeText.split("\n")
        html_lines = []
        if lines:
            first_line = lines[0].strip()
            if first_line.lower() == candName.lower():
                html_lines.append(f"<h2>{first_line}</h2>")
                lines.pop(0)
            else:
                html_lines.append(f"<h2>{candName}</h2>")
        for line in lines:
            s = line.strip()
            if not s:
                html_lines.append("<br>")
            elif s.endswith(":"):
                html_lines.append(f"<h3>{s}</h3>")
            else:
                html_lines.append(f"<p>{s}</p>")
        final_html = "\n".join(html_lines)
        self.text_edit.setStyleSheet("QTextEdit { font-family: 'Segoe UI', sans-serif; font-size: 11pt; margin: 8px; }")
        self.text_edit.setHtml(final_html)
        lay.addWidget(self.text_edit)
        btn_close = QtWidgets.QPushButton("Close")
        btn_close.clicked.connect(self.accept)
        lay.addWidget(btn_close)

class MainWindow(QtWidgets.QMainWindow):
    """Main window with tabs for manual entry, export/integrations, and training data."""
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Automated Search Strategy Generator")
        self.resize(1200, 800)
        self.setWindowIcon(QtGui.QIcon("favicon.ico"))
        self.statusBar().showMessage("Ready", 3000)
        central_widg = FuturisticWidget()
        self.setCentralWidget(central_widg)
        self.mainLayout = QtWidgets.QVBoxLayout(central_widg)
        self.mainLayout.setContentsMargins(10, 10, 10, 10)
        self.uiContainer = QtWidgets.QWidget()
        self.uiContainer.setAttribute(QtCore.Qt.WA_TranslucentBackground)
        self.mainLayout.addWidget(self.uiContainer)
        self.uiLayout = QtWidgets.QVBoxLayout(self.uiContainer)
        self.uiLayout.setSpacing(10)
        self.tabs = QtWidgets.QTabWidget()
        self.uiLayout.addWidget(self.tabs)

        # Create tabs
        self.tab_manual = QtWidgets.QWidget()
        self.tabs.addTab(self.tab_manual, "Manual Entry")
        self.init_manual_tab()
        self.tab_export = QtWidgets.QWidget()
        self.tabs.addTab(self.tab_export, "Export / Integrations")
        self.init_export_tab()
        self.tab_training = QtWidgets.QWidget()
        self.tabs.addTab(self.tab_training, "Training Data")
        self.init_training_tab()

        self.generator = FineTunedSearchStrategyGenerator("fine-tuned-t5")
        self.candidate_data = []  # Loaded from candidate_resumes.csv
        self.headers_candidate = []
        self.displayed_candidates = []

    def init_manual_tab(self):
        lay = QtWidgets.QVBoxLayout(self.tab_manual)
        lbl = QtWidgets.QLabel("Enter Job Description:")
        lay.addWidget(lbl)
        self.text_input = QtWidgets.QTextEdit()
        self.text_input.setPlaceholderText("Paste or type the job description here...")
        lay.addWidget(self.text_input)
        self.btn_generate = QtWidgets.QPushButton("Generate Boolean Query")
        self.btn_generate.clicked.connect(self.generate_query)
        lay.addWidget(self.btn_generate)
        lbl_q = QtWidgets.QLabel("Generated Boolean Query (Editable):")
        lay.addWidget(lbl_q)
        self.text_query = QtWidgets.QTextEdit()
        self.text_query.setPlaceholderText("The generated Boolean query will appear here. Feel free to tweak it!")
        lay.addWidget(self.text_query)

    def generate_query(self):
        jobDesc = self.text_input.toPlainText().strip()
        if not jobDesc:
            self.statusBar().showMessage("Please enter a job description.", 5000)
            return
        finalQuery, _, _ = self.generator.generate_boolean_query(jobDesc)
        self.text_query.setPlainText(finalQuery)

    def init_export_tab(self):
        lay = QtWidgets.QVBoxLayout(self.tab_export)
        group_export = QtWidgets.QGroupBox("Export / Integrations")
        groupLay = QtWidgets.QVBoxLayout(group_export)
        lbl = QtWidgets.QLabel("Export the generated query or open an external search on LinkedIn.\nAlso, search your local candidate dataset (candidate_resumes.csv).")
        lbl.setWordWrap(True)
        groupLay.addWidget(lbl)
        formLay = QtWidgets.QFormLayout()
        self.combo_format = QtWidgets.QComboBox()
        self.combo_format.addItems(["CSV", "JSON"])
        formLay.addRow("Output Format:", self.combo_format)
        groupLay.addLayout(formLay)
        self.btn_export = QtWidgets.QPushButton("Export Query")
        self.btn_export.clicked.connect(self.export_query)
        groupLay.addWidget(self.btn_export)
        group_plat = QtWidgets.QGroupBox("Open Query in External Platform")
        platLay = QtWidgets.QHBoxLayout(group_plat)
        btn_link = QtWidgets.QPushButton("LinkedIn")
        btn_link.clicked.connect(self.open_linkedin_search)
        platLay.addWidget(btn_link)
        groupLay.addWidget(group_plat)
        lay.addWidget(group_export)
        self.btn_search_candidates = QtWidgets.QPushButton("Search Local Candidates")
        self.btn_search_candidates.clicked.connect(self.search_local_candidates)
        lay.addWidget(self.btn_search_candidates)
        self.table_candidates = QtWidgets.QTableWidget()
        self.table_candidates.setSortingEnabled(True)
        self.table_candidates.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.table_candidates.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)
        lay.addWidget(self.table_candidates)
        btnLay = QtWidgets.QHBoxLayout()
        self.btn_view_resume = QtWidgets.QPushButton("View Resume")
        self.btn_view_resume.clicked.connect(self.view_resume)
        btnLay.addWidget(self.btn_view_resume)
        self.btn_export_resume = QtWidgets.QPushButton("Export Resume")
        self.btn_export_resume.clicked.connect(self.export_resume)
        btnLay.addWidget(self.btn_export_resume)
        lay.addLayout(btnLay)
        self.tab_export.setLayout(lay)

    def export_query(self):
        query = self.text_query.toPlainText().strip()
        if not query or query == "No keywords extracted.":
            self.statusBar().showMessage("Please generate a valid query before exporting.", 5000)
            return
        fmt = self.combo_format.currentText()
        filters = "CSV Files (*.csv);;JSON Files (*.json)"
        default_filter = "CSV Files (*.csv)" if fmt == "CSV" else "JSON Files (*.json)"
        opts = QtWidgets.QFileDialog.Options()
        fileName, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Query", "", filters, default_filter, options=opts)
        if fileName:
            try:
                if fmt == "CSV" or fileName.endswith(".csv"):
                    with open(fileName, mode="w", newline="", encoding="utf-8") as f:
                        writer = csv.writer(f)
                        writer.writerow(["Boolean Query"])
                        writer.writerow([query])
                    self.statusBar().showMessage(f"Query exported to: {fileName}", 5000)
                elif fmt == "JSON" or fileName.endswith(".json"):
                    with open(fileName, mode="w", encoding="utf-8") as f:
                        json.dump({"boolean_query": query}, f, indent=4)
                    self.statusBar().showMessage(f"Query exported to: {fileName}", 5000)
                else:
                    self.statusBar().showMessage("Unrecognized file format.", 5000)
            except Exception as ex:
                self.statusBar().showMessage(f"Export failed: {str(ex)}", 5000)

    def open_linkedin_search(self):
        query = self.text_query.toPlainText().strip()
        if not query or query == "No keywords extracted.":
            self.statusBar().showMessage("Generate a valid query first.", 5000)
            return
        simple = flatten_for_platform(query)
        encoded = quote_plus(simple)
        url = f"https://www.linkedin.com/search/results/people/?keywords={encoded}"
        webbrowser.open(url)

    def search_local_candidates(self):
        query = self.text_query.toPlainText().strip()
        if not query or query == "No keywords extracted.":
            self.statusBar().showMessage("Generate a valid query first.", 5000)
            return
        if not self.candidate_data:
            cand_file = "candidate_resumes.csv"
            if not os.path.exists(cand_file):
                self.statusBar().showMessage(f"Place {cand_file} in the root folder.", 5000)
                return
            try:
                df = pd.read_csv(cand_file)
                self.candidate_data = df.to_dict(orient="records")
                self.headers_candidate = list(df.columns)
                self.statusBar().showMessage(f"Loaded {len(self.candidate_data)} candidate resumes.", 5000)
            except Exception as e:
                self.statusBar().showMessage(str(e), 5000)
                return
        all_scores = []
        jobDesc = self.text_input.toPlainText().strip()
        for row in self.candidate_data:
            resumeTxt = row.get("Full Resume", row.get("Resume", row.get("resume", "")))
            # Use lower-case query for Boolean search
            _, score = hybrid_search_t5(resumeTxt, query.lower(), jobDesc,
                                        self.generator.model, self.generator.tokenizer,
                                        alpha=0.9, beta=0.1, thresh=0.0)
            all_scores.append((row, score))
        all_scores.sort(key=lambda x: x[1], reverse=True)
        if not all_scores:
            self.table_candidates.setRowCount(0)
            self.table_candidates.setColumnCount(0)
            self.statusBar().showMessage("No candidates found.", 5000)
            return
        high_rel = [(r, s) for (r, s) in all_scores if s >= 0.80]
        if high_rel:
            candidates = [r for (r, s) in high_rel]
            scores = high_rel
            msg = f"Found {len(high_rel)} high relevance candidate(s)."
        else:
            candidates = [r for (r, s) in all_scores[:10]]
            scores = all_scores[:10]
            msg = f"Showing top {len(candidates)} candidate(s)."
        self.populate_candidate_table(candidates, scores)
        self.statusBar().showMessage(msg, 5000)

    def populate_candidate_table(self, candidates, scores):
        if not candidates:
            self.table_candidates.setRowCount(0)
            self.table_candidates.setColumnCount(0)
            return
        headers = self.headers_candidate + ["Score"]
        self.table_candidates.setColumnCount(len(headers))
        self.table_candidates.setHorizontalHeaderLabels(headers)
        self.table_candidates.setRowCount(len(candidates))
        self.displayed_candidates = candidates
        for i, row in enumerate(candidates):
            for j, head in enumerate(self.headers_candidate):
                self.table_candidates.setItem(i, j, QtWidgets.QTableWidgetItem(str(row.get(head, ""))))
            sc = next(s for (r, s) in scores if r == row)
            self.table_candidates.setItem(i, len(headers) - 1, QtWidgets.QTableWidgetItem(f"{sc*100:.0f}%"))

    def view_resume(self):
        selected = self.table_candidates.selectionModel().selectedRows()
        if not selected:
            self.statusBar().showMessage("Select a candidate to view resume.", 5000)
            return
        candidate = self.displayed_candidates[selected[0].row()]
        resumeTxt = candidate.get("Full Resume", candidate.get("Resume", candidate.get("resume", "No resume available.")))
        dlg = ResumeDialog(candidate.get("Name", "Candidate"), resumeTxt)
        dlg.exec_()

    def export_resume(self):
        selected = self.table_candidates.selectionModel().selectedRows()
        if not selected:
            self.statusBar().showMessage("Select a candidate to export resume.", 5000)
            return
        candidate = self.displayed_candidates[selected[0].row()]
        resumeTxt = candidate.get("Full Resume", candidate.get("Resume", candidate.get("resume", "No resume available.")))
        candName = candidate.get("Name", "Candidate")
        opts = QtWidgets.QFileDialog.Options()
        fName, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Resume as DOCX", f"{candName}.docx", "DOCX Files (*.docx)", options=opts)
        if not fName:
            return
        try:
            doc = Document()
            # Setting style manually—this is a hack but it works
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            para_fmt = style.paragraph_format
            para_fmt.line_spacing = 1.15
            para_fmt.space_after = Pt(6)
            para_fmt.space_before = Pt(0)
            lines = resumeTxt.split("\n")
            processed = []
            for line in lines:
                if not line.strip() and processed and not processed[-1].strip():
                    continue
                processed.append(line)
            if processed:
                first = processed[0].strip()
                if first.lower() == candName.lower():
                    doc.add_heading(first, level=1)
                    processed.pop(0)
                else:
                    doc.add_heading(candName, level=1)
            for line in processed:
                s = line.strip()
                if not s:
                    doc.add_paragraph("")
                elif s.endswith(":"):
                    doc.add_heading(s, level=2)
                else:
                    doc.add_paragraph(s)
            doc.save(fName)
            self.statusBar().showMessage(f"Resume exported to: {fName}", 5000)
        except Exception as ex:
            self.statusBar().showMessage(f"Export failed: {str(ex)}", 5000)

    def init_training_tab(self):
        lay = QtWidgets.QVBoxLayout(self.tab_training)
        lbl = QtWidgets.QLabel("Training Data Preview (from Cleaned_Dataset.csv).\nPlace Cleaned_Dataset.csv in the root folder to view training examples.")
        lbl.setWordWrap(True)
        lay.addWidget(lbl)
        filtLay = QtWidgets.QHBoxLayout()
        lbl_filt = QtWidgets.QLabel("Filter (enter keyword):")
        filtLay.addWidget(lbl_filt)
        self.line_filter = QtWidgets.QLineEdit()
        self.line_filter.setPlaceholderText("Type to filter rows...")
        self.line_filter.textChanged.connect(self.filter_training_data)
        filtLay.addWidget(self.line_filter)
        lay.addLayout(filtLay)
        self.table_training = QtWidgets.QTableWidget()
        self.table_training.setSortingEnabled(True)
        self.table_training.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.table_training.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)
        lay.addWidget(self.table_training)
        self.btn_reload_data = QtWidgets.QPushButton("Reload Cleaned_Dataset.csv")
        self.btn_reload_data.clicked.connect(self.load_training_data)
        lay.addWidget(self.btn_reload_data)
        self.text_query_training = QtWidgets.QTextEdit()
        self.text_query_training.setReadOnly(True)
        self.text_query_training.setPlaceholderText("Generated query for the selected row will appear here.")
        lay.addWidget(self.text_query_training)
        self.btn_show_query = QtWidgets.QPushButton("Show Query for Selected Row")
        self.btn_show_query.clicked.connect(self.show_query_from_row)
        lay.addWidget(self.btn_show_query)
        self.training_data = []
        self.load_training_data()
        self.tab_training.setLayout(lay)

    def load_training_data(self):
        f_path = "Cleaned_Dataset.csv"
        if not os.path.exists(f_path):
            self.statusBar().showMessage(f"Could not find {f_path}.", 5000)
            return
        try:
            with open(f_path, mode="r", encoding="utf-8") as f:
                data = list(csv.reader(f))
            if not data:
                self.statusBar().showMessage(f"{f_path} is empty.", 5000)
                return
            hdrs = data[0]
            rows = data[1:]
            self.training_data = [dict(zip(hdrs, row)) for row in rows]
            self.populate_training_table(self.training_data)
            self.statusBar().showMessage(f"Loaded {len(rows)} rows from {f_path}.", 5000)
        except Exception as e:
            self.statusBar().showMessage(str(e), 5000)

    def populate_training_table(self, data):
        if not data:
            self.table_training.setRowCount(0)
            self.table_training.setColumnCount(0)
            return
        hdrs = list(data[0].keys())
        self.table_training.setColumnCount(len(hdrs))
        self.table_training.setHorizontalHeaderLabels(hdrs)
        self.table_training.setRowCount(len(data))
        for i, row in enumerate(data):
            for j, head in enumerate(hdrs):
                self.table_training.setItem(i, j, QtWidgets.QTableWidgetItem(row[head]))

    def filter_training_data(self):
        filt_text = self.line_filter.text().lower()
        filtered = self.training_data if not filt_text else [row for row in self.training_data if any(filt_text in str(v).lower() for v in row.values())]
        self.populate_training_table(filtered)

    def show_query_from_row(self):
        selected = self.table_training.selectionModel().selectedRows()
        if not selected:
            self.statusBar().showMessage("Select a row in the training data table.", 5000)
            return
        row = selected[0].row()
        job_desc = self.training_data[row].get("job_description", "")
        if not job_desc:
            self.statusBar().showMessage("Selected row has no job description.", 5000)
            return
        final_q, _, _ = self.generator.generate_boolean_query(job_desc)
        self.text_query_training.setPlainText(final_q)

    def show_message(self, title, txt, icon=QtWidgets.QMessageBox.Information):
        self.statusBar().showMessage(txt, 5000)

def main():
    app = QtWidgets.QApplication(sys.argv)
    app.setWindowIcon(QtGui.QIcon("favicon.ico"))
    # I love dark themes—if it's too gloomy, adjust as needed.
    styleSheet = """
    QMainWindow { background-color: #0a0a1e; }
    QWidget { color: #ffffff; font-family: 'Segoe UI', sans-serif; font-size: 12pt; }
    QTabWidget::pane { border: 1px solid #444; background: transparent; }
    QTabBar::tab { background: #1e1e2e; color: #ffffff; padding: 10px; margin: 2px; border-radius: 4px; }
    QTabBar::tab:selected { background: #3a3a50; color: #00ffff; }
    QPushButton {
        background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:0,
            stop:0 #3a3a50, stop:1 #00ffff);
        border: none; color: #ffffff; padding: 10px 20px; border-radius: 5px;
    }
    QPushButton:hover {
        background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:0,
            stop:0 #00ffff, stop:1 #3a3a50);
    }
    QTextEdit, QTableWidget, QComboBox, QGroupBox, QLineEdit {
        background-color: #1e1e2e;
        color: #ffffff;
        border: 1px solid #444;
        padding: 5px;
        border-radius: 4px;
    }
    QGroupBox {
        margin-top: 10px;
        font-weight: bold;
        border: 1px solid #444;
        border-radius: 6px;
        padding: 10px;
    }
    QGroupBox::title {
        subcontrol-origin: margin;
        subcontrol-position: top left;
        padding: 0 3px;
    }
    QHeaderView::section {
        background-color: #3a3a50;
        color: #00ffff;
        padding: 4px;
        border: none;
    }
    QDoubleSpinBox {
        background-color: #1e1e2e;
        color: #ffffff;
        border: 1px solid #444;
        padding: 2px;
        border-radius: 3px;
    }
    QComboBox QAbstractItemView {
        background-color: #1e1e2e;
        color: #ffffff;
        selection-background-color: #3a3a50;
        border: 1px solid #444;
        outline: 0;
    }
    QComboBox::drop-down {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 20px;
        border-left: 1px solid #444;
    }
    QMessageBox { background-color: #1e1e2e; color: #ffffff; }
    """
    app.setStyleSheet(styleSheet)
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    # If 'train' is passed as an argument, run training; otherwise, run the GUI.
    if len(sys.argv) > 1 and sys.argv[1] == "train":
        train_model()
    else:
        main()
